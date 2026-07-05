from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path
from typing import Any

import pandas as pd

from src.models import Document


MAX_CELL_LENGTH = 1200
CHUNK_WORDS = 260
CHUNK_OVERLAP = 45
SUPPORTED_EXTENSIONS = {
    ".csv",
    ".docx",
    ".jpeg",
    ".jpg",
    ".pdf",
    ".png",
    ".txt",
    ".xls",
    ".xlsx",
}


def normalize_column_name(name: Any) -> str:
    text = str(name).strip().lower()
    text = re.sub(r"[^a-z0-9]+", "_", text)
    return text.strip("_") or "column"


def clean_cell(value: Any) -> str:
    if pd.isna(value):
        return ""
    text = str(value).replace("\x00", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text[:MAX_CELL_LENGTH]


def dataframe_to_documents(df: pd.DataFrame, source_name: str) -> list[Document]:
    if df.empty:
        return []

    normalized = df.copy()
    normalized.columns = [normalize_column_name(col) for col in normalized.columns]

    docs: list[Document] = []
    for row_index, row in normalized.iterrows():
        parts = []
        metadata: dict[str, Any] = {
            "source_name": source_name,
            "row_number": int(row_index) + 2,
        }

        for column, value in row.items():
            cleaned = clean_cell(value)
            if not cleaned:
                continue
            parts.append(f"{column}: {cleaned}")
            if column in {"category", "section", "title", "therapy_area", "company", "product"}:
                metadata[column] = cleaned

        text = "\n".join(parts).strip()
        if not text:
            continue

        digest = hashlib.sha256(f"{source_name}:{row_index}:{text}".encode("utf-8")).hexdigest()
        docs.append(Document(id=digest, text=text, metadata=metadata))

    return docs


def file_to_documents(filename: str, data: bytes, mime_type: str = "") -> list[Document]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        return []

    if extension in {".csv", ".xls", ".xlsx"}:
        return tabular_file_to_documents(filename, data, extension)

    if extension == ".pdf":
        pages = extract_pdf_pages(data)
        return pages_to_documents(filename, pages, "pdf")

    if extension == ".docx":
        text = extract_docx_text(data)
        return text_to_documents(filename, text, "docx")

    if extension in {".png", ".jpg", ".jpeg"}:
        text = extract_image_text(data)
        return text_to_documents(filename, text, "image")

    text = data.decode("utf-8", errors="ignore")
    return text_to_documents(filename, text, mime_type or "text")


def tabular_file_to_documents(filename: str, data: bytes, extension: str) -> list[Document]:
    buffer = io.BytesIO(data)
    if extension == ".csv":
        try:
            df = pd.read_csv(buffer)
        except UnicodeDecodeError:
            buffer.seek(0)
            df = pd.read_csv(buffer, encoding="latin-1")
    else:
        df = pd.read_excel(buffer)
    return dataframe_to_documents(df, source_name=filename)


def extract_pdf_pages(data: bytes) -> list[tuple[int, str]]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append((page_number, text))
    return pages


def extract_docx_text(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def extract_image_text(data: bytes) -> str:
    from PIL import Image
    import pytesseract

    image = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(image)


def pages_to_documents(
    source_name: str,
    pages: list[tuple[int, str]],
    file_type: str,
) -> list[Document]:
    docs = []
    for page_number, page_text in pages:
        docs.extend(
            text_to_documents(
                source_name=source_name,
                text=page_text,
                file_type=file_type,
                base_metadata={"page_number": page_number},
            )
        )
    return docs


def text_to_documents(
    source_name: str,
    text: str,
    file_type: str,
    base_metadata: dict[str, Any] | None = None,
) -> list[Document]:
    cleaned = re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()
    if not cleaned:
        return []

    words = cleaned.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + CHUNK_WORDS, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start = max(end - CHUNK_OVERLAP, start + 1)

    docs = []
    for chunk_index, chunk in enumerate(chunks, start=1):
        metadata: dict[str, Any] = {
            "source_name": source_name,
            "file_type": file_type,
            "chunk_number": chunk_index,
        }
        if base_metadata:
            metadata.update(base_metadata)
        digest = hashlib.sha256(
            f"{source_name}:{file_type}:{chunk_index}:{chunk}".encode("utf-8")
        ).hexdigest()
        docs.append(Document(id=digest, text=chunk, metadata=metadata))
    return docs
