"""Document parsing — PDF (PyMuPDF), DOCX, TXT. Returns per-page text.

Shared by the server delta-sync and the Mac ingestion CLI.
"""
import io
import re


def parse_file(path_or_bytes, mime_type: str, filename: str = "") -> list[str]:
    """Return list of page texts. Non-paginated formats return a single 'page'."""
    data = path_or_bytes
    if isinstance(path_or_bytes, str):
        with open(path_or_bytes, "rb") as f:
            data = f.read()

    name = filename.lower()
    if mime_type == "application/pdf" or name.endswith(".pdf"):
        return _parse_pdf(data)
    if "wordprocessingml" in mime_type or name.endswith(".docx"):
        return _parse_docx(data)
    if mime_type.startswith("text/") or name.endswith((".txt", ".md")):
        return [data.decode("utf-8", errors="replace")]
    raise ValueError(f"Unsupported type: {mime_type or filename}")


def _parse_pdf(data: bytes) -> list[str]:
    import fitz  # PyMuPDF
    pages = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            pages.append(_clean(page.get_text("text")))
    return pages


def _parse_docx(data: bytes) -> list[str]:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return [_clean("\n".join(parts))]


def _clean(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def looks_scanned(pages: list[str]) -> bool:
    """Heuristic: mostly-empty text layer means the PDF is image-only (needs OCR)."""
    if not pages:
        return True
    non_empty = sum(1 for p in pages if len(p) > 50)
    return non_empty / len(pages) < 0.2


YEAR_RE = re.compile(r"\b(19[89]\d|20[0-4]\d)\b")
AGENCIES = ["FDA", "EMA", "ICH", "WHO", "PIC/S", "PICS", "MHRA", "CDSCO", "TGA", "PMDA", "ANVISA", "USP"]


def guess_metadata(filename: str, first_page: str, default_agency: str = "") -> dict:
    """Best-effort agency/year extraction from filename + first page."""
    hay = f"{filename}\n{first_page[:2000]}".upper()
    agency = default_agency
    if not agency:
        for a in AGENCIES:
            if a in hay:
                agency = "PIC/S" if a == "PICS" else a
                break
    years = YEAR_RE.findall(filename) or YEAR_RE.findall(first_page[:2000])
    year = int(years[0]) if years else 0
    return {"agency": agency, "year": year}
